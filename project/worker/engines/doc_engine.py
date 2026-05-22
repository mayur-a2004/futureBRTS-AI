from docx import Document
import os

def generate_doc(payload):
    os.makedirs("outputs", exist_ok=True)
    doc = Document()
    doc.add_heading("Future BRTS Document", level=1)
    doc.add_paragraph(payload.prompt or "")

    path = f"outputs/doc_{payload.taskId}.docx"
    doc.save(path)

    return {"status": "SUCCESS", "artifacts": {"doc": path}}
