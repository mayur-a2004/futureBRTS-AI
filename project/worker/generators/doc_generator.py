import os
import json
import re
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.diagram_engine import diagram_engine
from utils.sanitizer import sanitize_for_xml

PROJECTS_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "backend", "projects"))
os.makedirs(PROJECTS_ROOT, exist_ok=True)

def sanitize_for_xml(text):
    """Removes control characters that break XML-based formats (DOCX/PPTX)."""
    if not text: return ""
    illegal_chars = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f\ud800-\udfff\ufdd0-\ufdef\ufffe\uffff]')
    return illegal_chars.sub('', str(text))

def _set_font(paragraph, name='Times New Roman', size=12, bold=False):
    for run in paragraph.runs:
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold

def generate_project_doc(project_name, category, field, project_type, tech_stack, vision, ai_content, job_id, images=None):
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin, section.bottom_margin = Cm(2.5), Cm(2.5)
        section.left_margin, section.right_margin = Cm(3.0), Cm(2.0)

    # Load Master Blueprint (Academic Standard) - Fix 5: Absolute Path
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    blueprint_path = os.path.join(base_dir, "storage", "templates", "master_blueprint.json")
    master_blueprint = {}
    if os.path.exists(blueprint_path):
        with open(blueprint_path, "r") as f:
            master_blueprint = json.load(f)
    else:
        print(f"⚠️ MASTER BLUEPRINT NOT FOUND AT {blueprint_path}")

    # Helper for AI extraction (More robust)
    def get_section_text(num, title, fallback):
        # Look for Chapter Headers or Section Numbers
        pattern = rf"(?:{num}\.?\s+|[Cc]hapter\s*{num}\.?\s+){title}(.*?)(?=(?:\d+\.?\s+|[Cc]hapter\s*\d+\.?\s+|$))"
        match = re.search(pattern, ai_content, re.DOTALL | re.IGNORECASE)
        
        extracted_content = fallback
        mermaid_codes = []

        if match:
            extracted_content = match.group(1).strip().replace("*", "")
        
        # SEARCH FOR MERMAID BLOCKS
        # Pattern for [MERMAID_START]...[MERMAID_END]
        mermaid_pattern = r"\[MERMAID_START\](.*?)\[MERMAID_END\]"
        matches = re.findall(mermaid_pattern, ai_content if not match else match.group(1), re.DOTALL)
        for m in matches:
            mermaid_codes.append(m.strip())
            
        return extracted_content, mermaid_codes

    # 1. COVER PAGE (Supreme Branding)
    doc.add_paragraph("\n")
    if images and 'logo' in images:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(images['logo'], width=Inches(2.0))
        except: pass

    doc.add_paragraph("\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"TITAN INDUSTRIAL DOSSIER\n{project_name.upper()}")
    r.font.size, r.font.bold, r.font.name = Pt(28), True, 'Arial'
    
    doc.add_paragraph("\n" * 2)
    p = doc.add_paragraph(f"A Technical Project Report Submitted for the requirements of:\n{category}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p, size=14, bold=True)

    p = doc.add_paragraph(f"Domain: {field}\nArchitecture: {tech_stack}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p, size=12)

    doc.add_paragraph("\n" * 4)
    p = doc.add_paragraph(f"GENESIS NEURAL ARCHITECT v5.0\nIndustrial Compliance Verified\nYear: {datetime.datetime.now().year}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p, name='Arial', size=10)
    
    doc.add_page_break()
    
    # 2. INDUSTRIAL BOILERPLATE ENGINE (RECORDS MATCHING 80-PAGE STANDARD)
    academic_boilerplate = {
        "Abstract": f"This project, {project_name}, represents a strategic implementation of {tech_stack} to solve critical industrial gaps in the {field} domain. The system leverages a 3-tier architecture, ensuring scalability, security, and real-time data processing. By integrating high-fidelity neural logic, the solution provides an optimized roadmap for {vision}.",
        "Introduction": f"The evolution of {field} has necessitated highly responsive and intelligent systems. This report details the architectural genesis of {project_name}, focusing on industrial workflows and technical sustainability. The implementation follows standard SDLC methodologies to ensure production-grade reliability.",
        "System Requirements": "HARDWARE: Minimum 8GB RAM, Quad-core Processor, 256GB SSD Architecture.\nSOFTWARE: Modern Browser Support (Chrome/Edge), Node.js Runtime, Database Management System, and secure SSL/TLS protocols.",
        "Feasibility Study": "TECHNICAL FEASIBILITY: The chosen stack is industry-standard and highly compatible with modern cloud environments. ECONOMIC FEASIBILITY: The solution offers high ROI through automated workflows. OPERATIONAL FEASIBILITY: Strategic UI design ensures minimal training requirements for end-users.",
        "System Design": "The system utilizes a Decoupled 3-Tier Architecture. LAYER 1: Presentation Layer (React/Vite) providing a rich visual interface. LAYER 2: Logic Layer (Express/FastAPI) managing business rules and authentication. LAYER 3: Data Layer (MongoDB/PostgreSQL) ensuring 3NF ACID compliance.",
        "Database Design": f"The database schema for {project_name} is designed for 3rd Normal Form (3NF) efficiency. It includes Entity-Relationship mapping for Users, Transactions, and System Analytics, ensuring data integrity across large-scale industrial operations.",
        "Testing": f"The testing phase for {project_name} includes Unit Testing for core components, Integration Testing for API bridges, and User Acceptance Testing (UAT) to validate the vision of {vision}. All test cases have been verified against industrial benchmarks.",
        "Future Scope": "The roadmap includes integration of Machine Learning for predictive analytics, cross-platform mobile deployment, and transition to a Microservices-based distributed architecture for global scalability."
    }

    sections_to_render = master_blueprint.get("structure", [])
    if not sections_to_render:
        sections_to_render = [
            {"text": "Certificate", "fallback": "This is to certify that the project work titled '{name}' is a bonafide work carried out in partial fulfillment of the academic requirements.".format(name=project_name)},
            {"text": "Declaration", "fallback": "I hereby declare that this project report is my original work and has not been submitted to any other university or institution."},
            {"text": "Acknowledgement", "fallback": "I would like to express my gratitude to the FutureBRTS Neural Core and my mentors for their guidance during this industrial project implementation."},
            {"text": "Abstract", "fallback": academic_boilerplate["Abstract"]},
            {"text": "Introduction", "fallback": academic_boilerplate["Introduction"]},
            {"text": "Problem Statement", "fallback": "Current systems lack the required scalability and high-fidelity synchronization needed for modern industrial demands. This project addresses these gaps through {stack}.".format(stack=tech_stack)},
            {"text": "Objectives", "fallback": "1. Implement a secure 3-tier architecture.\n2. Ensure 99.9% data availability.\n3. Provide a responsive industrial dashboard UI."},
            {"text": "Scope", "fallback": "The scope covers full-stack development, including secure API orchestration, database normalization, and responsive visual implementation for global users."},
            {"text": "System Requirements", "fallback": academic_boilerplate["System Requirements"]},
            {"text": "Feasibility Study", "fallback": academic_boilerplate["Feasibility Study"]},
            {"text": "SRS", "fallback": "The Software Requirement Specification follows IEEE-830 standards, detailing functional and non-functional requirements for high-availability systems."},
            {"text": "System Design", "fallback": academic_boilerplate["System Design"]},
            {"text": "Database Design", "fallback": academic_boilerplate["Database Design"]},
            {"text": "Data Flow Diagrams", "fallback": "Level 0, Level 1, and Level 2 DFDs showcase the transactional flow from user initiation to database persistence and visual feedback."},
            {"text": "Module Description", "fallback": "The system is divided into three core modules: Authentication Engine (Secure Access), Project Logic (Business Rules), and Admin Analytics (Strategic Insights)."},
            {"text": "Implementation", "fallback": "The implementation phase involved setting up the {stack} environment, configuring secure API middleware, and deploying responsive front-end components.".format(stack=tech_stack)},
            {"text": "Source Code Strategy", "fallback": "Logic follows Clean Code principles, focusing on modularity, high cohesion, and low coupling to ensure maintainable industrial software."},
            {"text": "Testing", "fallback": academic_boilerplate["Testing"]},
            {"text": "Visual Representation", "fallback": "The user interface is designed with a focus on 'Information Density' and 'Visual Hierarchy', utilizing modern CSS frameworks for responsiveness."},
            {"text": "Results", "fallback": "System performance benchmarks indicate a high response rate and successful resolution of all primary industrial problem statements."},
            {"text": "Future Scope", "fallback": academic_boilerplate["Future Scope"]},
            {"text": "References", "fallback": "1. Official Documentation for {stack}\n2. IEEE Industrial Software Standards\n3. Modern Web Architecture (Academic Press)".format(stack=tech_stack)}
        ]

    for i, section in enumerate(sections_to_render, 1):
        title = section.get("text", f"Section {i}")
        fallback_msg = section.get("fallback", "Detailed technical synthesis provided in the industrial source package.")
        
        h = doc.add_heading(f"{i}. {title}", level=1)
        for run in h.runs:
            run.font.name, run.font.size, run.font.color.rgb = 'Arial', Pt(18), RGBColor(31, 73, 125)
        
        content_text, mermaid_diagrams = get_section_text(i, title, fallback_msg)
        p = doc.add_paragraph(sanitize_for_xml(content_text if len(content_text) > 10 else fallback_msg))
        _set_font(p, 'Times New Roman', 12)
        
        # RENDER AND EMBED DIAGRAMS
        diagram_dir = os.path.join("storage", "projects", job_id, "diagrams")
        os.makedirs(diagram_dir, exist_ok=True)
        
        # Embed existing pre-generated diagrams if any (Fix Engine Step 5/8/9)
        if i == 12 or i == 13 or i == 20: # Design/DB/Outcome sections
            for f in os.listdir(diagram_dir):
                if f.endswith(".png"):
                    try:
                        f_path = os.path.join(diagram_dir, f)
                        doc.add_picture(f_path, width=Inches(5.0))
                        doc.add_paragraph(f"Industrial Diagram: {f.replace('_',' ').replace('.png','')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except: pass

        for idx, m_code in enumerate(mermaid_diagrams):
             try:
                  img_path = diagram_engine.render_mermaid(m_code, output_dir=diagram_dir)
                  if img_path:
                       p_img = doc.add_paragraph()
                       p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                       doc.add_picture(img_path, width=Inches(5.0))
                       doc.add_paragraph(f"Figure {i}.{idx+1}: Industrial {title} Schematic").alignment = WD_ALIGN_PARAGRAPH.CENTER
             except: pass

        # Add visual context if matching UI section
        if ("Visual" in title or "Design" in title or i == 19) and images and 'mockup' in images:
            try:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_picture(images['mockup'], width=Inches(5.5))
            except: pass

        doc.add_page_break()

    safe_name = re.sub(r'[^\w\s-]', '', project_name).strip().replace(' ', '_')
    filename = f"Project_Report_{safe_name}_{job_id[:8]}.docx"
    doc_dir = os.path.join(PROJECTS_ROOT, job_id, "docs")
    os.makedirs(doc_dir, exist_ok=True)
    path = os.path.join(doc_dir, filename)
    doc.save(path)
    return path
